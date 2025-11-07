class DocUtils:
    """Utilidad para operaciones y manipulación básica de archivos DOCX."""

    @staticmethod
    def leer_docx(nombre_archivo):
        """Lee el texto de un archivo DOCX."""
        try:
            from docx import Document
            doc = Document(nombre_archivo)
            contenido = []
            for p in doc.paragraphs:
                contenido.append(p.text)
            return "\n".join(contenido)
        except Exception as e:
            print(f"Error leyendo DOCX: {e}")
            return ""

    @staticmethod
    def escribir_docx(nombre_archivo, lineas, titulo=None):
        """Escribe texto en un archivo DOCX, opcionalmente con título."""
        try:
            from docx import Document
            doc = Document()
            if titulo:
                doc.add_heading(titulo, level=1)
            for linea in lineas:
                doc.add_paragraph(linea)
            doc.save(nombre_archivo)
            return True
        except Exception as e:
            print(f"Error escribiendo DOCX: {e}")
            return False
